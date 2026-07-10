"""Novel upload document parsing helpers."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

TEXT_NOVEL_EXTENSIONS = {".txt", ".md"}
SUPPORTED_NOVEL_EXTENSIONS = TEXT_NOVEL_EXTENSIONS | {".docx"}
SUPPORTED_NOVEL_EXTENSION_ORDER = (".txt", ".md", ".docx")
_BILLABLE_WHITESPACE_RE = re.compile(r"[\s\u3000]+")


@dataclass
class DocumentParseError(ValueError):
    """Raised when an uploaded novel cannot be converted to plain text."""

    source_format: str
    location: str
    reason: str
    raw_exception: Exception | None = None

    def __str__(self) -> str:
        return f"{self.location}: {self.reason}"


def supported_novel_extensions_label() -> str:
    return "、".join(SUPPORTED_NOVEL_EXTENSION_ORDER)


def is_supported_novel_path(path: str | Path) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_NOVEL_EXTENSIONS


def count_billable_novel_chars(text: str) -> int:
    """Count parsed novel text characters used for import billing.

    The import pipeline works on decoded/extracted plain text, so billing uses
    that same text and ignores layout whitespace. Punctuation remains billable
    because it is still model-visible input.
    """
    if not text:
        return 0
    return len(_BILLABLE_WHITESPACE_RE.sub("", text))


def decode_novel_bytes(raw: bytes) -> str:
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("gbk")


def load_novel_text(path: str | Path) -> str:
    novel_path = Path(path)
    suffix = novel_path.suffix.lower()
    if suffix in TEXT_NOVEL_EXTENSIONS:
        try:
            text = decode_novel_bytes(novel_path.read_bytes())
            # CRLF/CR novels (typical Windows txt) must not leak \r downstream.
            return text.replace("\r\n", "\n").replace("\r", "\n")
        except UnicodeDecodeError as exc:
            raise DocumentParseError(
                source_format=suffix.lstrip(".") or "unknown",
                location="文件编码",
                reason="仅支持 UTF-8 或 GBK 编码的文本文件",
                raw_exception=exc,
            ) from exc

    if suffix == ".docx":
        return parse_docx(novel_path)

    raise DocumentParseError(
        source_format=suffix.lstrip(".") or "unknown",
        location="文件类型",
        reason=(
            f"不支持的文件类型: {suffix or '无扩展名'}，"
            f"当前支持: {supported_novel_extensions_label()}"
        ),
    )


def parse_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentParseError(
            source_format="docx",
            location="依赖",
            reason="缺少 python-docx，无法解析 Word 文档",
            raw_exception=exc,
        ) from exc

    try:
        document = Document(path)
    except Exception as exc:
        raise DocumentParseError(
            source_format="docx",
            location="文档",
            reason="无法打开 Word 文档",
            raw_exception=exc,
        ) from exc

    paragraphs: list[str] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        try:
            text = paragraph.text.strip()
        except Exception as exc:
            raise DocumentParseError(
                source_format="docx",
                location=f"段落 {index}",
                reason="读取段落失败",
                raw_exception=exc,
            ) from exc
        if text:
            paragraphs.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        try:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append("\t".join(cells))
        except Exception as exc:
            raise DocumentParseError(
                source_format="docx",
                location=f"表格 {table_index}",
                reason="读取表格失败",
                raw_exception=exc,
            ) from exc

    return "\n\n".join(paragraphs)
